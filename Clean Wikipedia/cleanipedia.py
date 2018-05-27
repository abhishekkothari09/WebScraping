import bs4
from bs4 import BeautifulSoup as soup
from urllib.request import urlopen as uOpen
from docx import Document

def html_parsing(my_url):

    # Open the connection and get the page.
    opn = uOpen(my_url)
    page_html = opn.read()
    opn.close()

    # HTML parsing
    page_soup = soup(page_html, 'html.parser')
    return page_soup


def remove_links(page_soup):

    #ptags = page_soup.findAll("div", {"class": "mw-parser-output"})
    #container = page_soup.findAll("div", {"class": "mw-body"})

    #for div in divs:
    #    h1 = div.find("h1")
    #    h2 = div.find("h2")
    #    h3 = div.find("h3")
    #    print(h1.text)
    #    print(h2.text)
    #    print(h3.text)




    ptags = page_soup.findAll("p")
    #print(ptags)
    for ptag in ptags:
        print(ptag.text)
        document.add_paragraph(ptag.text)

# First page
my_url = 'https://en.wikipedia.org/wiki/Film'
page_soup = html_parsing(my_url)
document = Document()
cleaned = remove_links(page_soup)
#document.add_paragraph(page_soup.text)
#print(page_soup)

document.save('demo.docx')